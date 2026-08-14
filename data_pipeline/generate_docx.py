#!/usr/bin/env python3
"""Generate a warning letter .docx from a template."""

from pathlib import Path

from docx import Document


def generate_warning_letter(
    output_path: Path,
    airline_name: str,
    flight_number: str,
    flight_date: str,
    origin: str,
    destination: str,
    delay_hours: str,
    amount: str,
    passenger_name: str,
    phone: str,
    email: str,
) -> Path:
    doc = Document()
    doc.add_heading("מכתב התראה / Warning Letter", level=1)
    doc.add_paragraph(f"To: {airline_name}")
    doc.add_paragraph(f"Re: Flight {flight_number} on {flight_date}")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph(
        f"I am writing regarding flight {flight_number} from {origin} to {destination} on {flight_date}. "
        f"The flight was delayed/cancelled by approximately {delay_hours} hours."
    )
    doc.add_paragraph(
        "Under the Israeli Aviation Services Law (Tibi Law) / EU Regulation 261, I am entitled to compensation."
    )
    doc.add_paragraph(f"I hereby demand payment of {amount} within 14 days of receipt of this letter.")
    doc.add_paragraph("If payment is not received, I intend to pursue legal action without further notice.")
    doc.add_paragraph("")
    doc.add_paragraph(f"Sincerely,\n{passenger_name}\n{phone}\n{email}")

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "templates" / "sample-warning-letter.docx"
    generate_warning_letter(
        output_path=out,
        airline_name="El Al",
        flight_number="LY315",
        flight_date="2024-06-12",
        origin="TLV",
        destination="JFK",
        delay_hours="11",
        amount="3,000 ILS",
        passenger_name="Passenger Name",
        phone="+972-50-000-0000",
        email="passenger@example.com",
    )
    print(f"Generated {out}")
